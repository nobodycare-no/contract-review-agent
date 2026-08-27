"""T5 工具面/状态机集成测试：mock 审批系统替换为进程内假件，全链路离线可回归。"""
from __future__ import annotations

import io

import docx
import pytest
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session


def _docx_bytes() -> bytes:
    document = docx.Document()
    document.add_heading("GPU 服务器集群采购合同", level=0)
    for para in (
        "合同编号：HT-2026-0301",
        "甲方（采购方）：XX科技有限公司（统一社会信用代码：91310000MA1FL8X20A）",
        "乙方（供应商）：华信计算设备有限公司",
        "合同总金额为人民币 1,860,000 元（大写：壹佰捌拾陆万元整），含税。",
        "预付款比例为合同总金额的 50%。",
        "剩余货款于到货后 90 个工作日内支付。",
        "因履行本合同发生争议的，任何一方可向甲方所在地人民法院提起诉讼。",
    ):
        document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


FAKE_ITEMS = [
    {"instance_id": "inst-001", "approval_code": "AP-X-001", "title": "采购合同审批",
     "applicant": "王铁柱", "apply_time": "2026-03-02 09:30"},
    {"instance_id": "inst-006", "approval_code": "AP-X-006", "title": "缺附件异常单",
     "applicant": "李梅", "apply_time": "2026-03-07 15:00"},
]
FAKE_DETAILS = {
    "inst-001": {"attachments": [{"attachment_id": "att-a", "file_name": "合同.docx"}]},
    "inst-006": {"attachments": []},
}


@pytest.fixture()
def fake_mock(monkeypatch):
    """进程内假外部系统：记录评论外呼次数以验证幂等守卫。"""
    posted: list[str] = []

    def fake_list(limit=20):
        return FAKE_ITEMS

    def fake_detail(instance_id):
        if instance_id not in FAKE_DETAILS:
            from app.services.tool_errors import ToolError
            raise ToolError("APPROVAL_NOT_FOUND", instance_id)
        return {**FAKE_DETAILS[instance_id], "instance_id": instance_id}

    def fake_download(instance_id, attachment_id):
        assert (instance_id, attachment_id) == ("inst-001", "att-a")
        return _docx_bytes(), "技术采购合同.docx"

    def fake_post(instance_id, comment_text):
        posted.append(f"{instance_id}:{len(posted) + 1}")
        return {"write_status": "success", "comment_id": len(posted)}

    import app.services.mock_client as mc

    monkeypatch.setattr(mc, "list_pending", fake_list)
    monkeypatch.setattr(mc, "get_detail", fake_detail)
    monkeypatch.setattr(mc, "download_attachment", fake_download)
    monkeypatch.setattr(mc, "post_comment", fake_post)
    return {"posted": posted}


def _post(client: TestClient, path: str, payload: dict) -> dict:
    resp = client.post(path, json=payload)
    assert resp.status_code == 200, resp.text
    return resp.json()


class TestFetchDedup:
    def test_second_pull_creates_nothing(self, client: TestClient, fake_mock) -> None:
        first = _post(client, "/tools/list_pending", {"limit": 10})
        assert first["data"]["sync"]["created"] == 2
        second = _post(client, "/tools/list_pending", {"limit": 10})
        assert second["data"]["sync"]["created"] == 0
        assert second["data"]["sync"]["total"] == 2  # 唯一业务标识去重：只更新不重建


class TestFullPipeline:
    def test_ac1_to_ac5_offline(self, client: TestClient, db_session: Session, fake_mock) -> None:
        from app.tools.bootstrap import seed_rules

        seed_rules(db_session)
        _post(client, "/tools/list_pending", {})
        tasks = client.get("/agent/tasks").json()["tasks"]
        buy = next(t for t in tasks if t["approval_code"] == "AP-X-001")

        # AC-2 下载
        downloaded = _post(client, "/tools/download_attachment",
                           {"instance_id": "inst-001"})
        assert downloaded["data"]["attachments"][0]["download_status"] == "done"

        # AC-3 解析
        parsed = _post(client, "/tools/parse_document", {"document_id": buy["id"]})
        assert parsed["data"]["task_status"] == "reviewing"
        assert parsed["data"]["basic_info"]["amount"]["value"] == 1_860_000.0

        # AC-4 规则
        ruled = _post(client, "/tools/run_rules", {"case_id": buy["id"]})
        codes = {h["rule_code"] for h in ruled["data"]["hits"]}
        assert {"PAY_ADVANCE_HIGH", "PAY_CYCLE_LONG", "NO_BREACH",
                "JURISDICTION_RISK"} <= codes
        assert ruled["data"]["overall_risk_level"] == "high"

        # AC-5 保存+回写
        saved = _post(client, "/tools/save_result", {
            "case_id": buy["id"], "overall_risk_level": ruled["data"]["overall_risk_level"],
            "summary_text": "摘要", "focus_points_json": ruled["data"]["focus_points"],
            "comment_text": reviewer_style(ruled)})
        review_id = saved["data"]["review_id"]
        written = _post(client, "/tools/write_comment",
                        {"instance_id": "inst-001", "review_id": review_id})
        assert written["data"]["write_status"] == "success"
        assert len(fake_mock["posted"]) == 1

        # 幂等守卫：重复回写不再外呼
        again = _post(client, "/tools/write_comment",
                      {"instance_id": "inst-001", "review_id": review_id})
        assert again["data"]["deduped"] is True
        assert len(fake_mock["posted"]) == 1

        detail = client.get(f"/agent/tasks/{buy['id']}").json()
        assert detail["task"]["task_status"] == "done"
        assert detail["task"]["write_status"] == "success"


def reviewer_style(ruled: dict) -> str:
    d = ruled["data"]
    label = {"high": "高", "medium": "中", "low": "低"}[d["overall_risk_level"]]
    lines = [f"【AI合同审查】总风险等级：{label}", "", "一、命中规则列表"]
    lines += [f"{i}. [{h['risk_level']}] {h['rule_name']}"
              for i, h in enumerate(d["hits"], 1)]
    lines += ["", "三、审批关注点"] + [f"- {p}" for p in d["focus_points"]]
    return "\n".join(lines)


class TestBlockedAndRetry:
    def test_ac6_block_then_retry(self, client: TestClient, db_session: Session, fake_mock) -> None:
        _post(client, "/tools/list_pending", {})
        tasks = client.get("/agent/tasks").json()["tasks"]
        broken = next(t for t in tasks if t["approval_code"] == "AP-X-006")

        result = _post(client, "/tools/parse_document", {"document_id": broken["id"]})
        assert result["ok"] is False
        assert result["error"]["code"] == "ATTACHMENT_MISSING"

        detail = client.get(f"/agent/tasks/{broken['id']}").json()
        assert detail["task"]["task_status"] == "blocked"
        assert "ATTACHMENT_MISSING" in detail["task"]["block_reason"]

        retried = client.post(f"/agent/tasks/{broken['id']}/retry")
        assert retried.status_code == 200
        assert retried.json()["resumed_stage"] == "parsing"


class TestGuards:
    def test_comment_missing_marker_auto_completed(self, client: TestClient, fake_mock) -> None:
        """G7 自动补齐语义：缺『总风险等级』行时护栏自动加标准行，
        模型首次产出欠规范零步数损耗；格式契约对下游恒成立。"""
        _post(client, "/tools/list_pending", {})
        tasks = client.get("/agent/tasks").json()["tasks"]
        buy = next(t for t in tasks if t["approval_code"] == "AP-X-001")
        saved = _post(client, "/tools/save_result", {
            "case_id": buy["id"], "overall_risk_level": "low",
            "summary_text": "s", "focus_points_json": [],
            "comment_text": "整体尚可，建议留意交付节点。"})
        assert saved["ok"] is True
        assert saved["data"]["comment_text"].startswith("【AI合同审查】总风险等级：低")
