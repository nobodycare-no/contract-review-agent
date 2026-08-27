"""V1 合并后共享测试工厂：真实本地审批域（不再 mock 外部 HTTP）。"""
from __future__ import annotations

import io

import docx


def docx_bytes() -> bytes:
    document = docx.Document()
    document.add_heading("GPU 服务器集群采购合同", level=0)
    for para in (
        "合同编号：HT-2026-0301",
        "甲方（采购方）：XX科技有限公司（统一社会信用代码：91310000MA1FL8X20A）",
        "乙方（供应商）：华信计算设备有限公司",
        "合同总金额为人民币 1,860,000 元（大写：壹佰捌拾陆万元整），含税。",
        "预付款比例为合同总金额的 50%。",
        "剩余货款于到货后 90 个工作日内支付。",
        "因履行本合同发生争议的，任何一方可向甲方所在地人民法院提起诉讼。",
    ):
        document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_form(db_session, *, code: str = "AP-Z-001", title: str = "采购合同审批",
              applicant: str = "王铁柱", files: int = 1):
    """经真实 create_form 建审批单（附件落盘 UPLOAD_DIR）；files=0 模拟缺附件单。"""
    if files == 0:
        """运行时防线测试专用：模拟绕过创建闸门的历史脏数据(零附件)。"""
        from app.models import ApprovalTask

        t = ApprovalTask(approval_code=code, approval_title=title,
                         applicant_name=applicant, instance_id=f"LOCAL-{code}")
        db_session.add(t); db_session.commit(); return t
    from app.services.approval_store import create_form

    sources = [("合同.docx", docx_bytes())] * files
    return create_form(title=title, applicant=applicant,
                       sources=sources, approval_code=code,
                       instance_id=f"LOCAL-{code}")


def post_spy(monkeypatch) -> list[str]:
    """包裹真实回写确认并捕获评论文本序列。"""
    import app.services.approval_store as store

    captured: list[str] = []
    original = store.post_comment

    def spy(instance_id, comment_text):
        captured.append(comment_text)
        return original(instance_id, comment_text)

    monkeypatch.setattr(store, "post_comment", spy)
    return captured
