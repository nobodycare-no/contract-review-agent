"""T3 解析器测试：归一化契约（金额数值元/日期锚定/币种推断/标题回退）+ docx 往返 + OCR 门控。"""
from __future__ import annotations

import shutil

import pytest

from app.services.parser import ParseError, extract_raw_text, extract_structured


def _basic(structured: dict, field: str) -> dict:
    return structured["basic_info"][field]


class TestNormalization:
    def test_amount_with_thousands_separator_to_yuan(self) -> None:
        text = "合同总金额为人民币 1,860,000 元（大写：壹佰捌拾陆万元整），含税。"
        amount = _basic(extract_structured(text), "amount")
        assert amount["status"] == "ok"
        assert amount["value"] == 1_860_000.0
        assert amount["raw_text"] == "1,860,000 元"
        assert isinstance(amount["pos"], int)

    def test_amount_wan_multiplier(self) -> None:
        text = "本合同总价 50 万元，含运输费用。"
        amount = _basic(extract_structured(text), "amount")
        assert amount["value"] == pytest.approx(500_000.0)
        assert amount["raw_text"] == "50 万元"  # 原文形态保真

    def test_amount_missing_status(self) -> None:
        structured = extract_structured("双方权利义务如下……")
        assert _basic(structured, "amount")["status"] == "missing"

    def test_date_anchor_avoids_false_positive(self) -> None:
        """『自营』等非日期语境不得触发生效日期；年份锚定生效。"""
        text = "公司自营产品线2026年扩张计划……"
        assert _basic(extract_structured(text), "effective_date")["status"] == "missing"

    def test_date_labeled_and_normalized(self) -> None:
        text = "生效日期：2026年3月1日；到期日期：2027年2月28日。"
        structured = extract_structured(text)
        assert _basic(structured, "effective_date")["value"] == "2026-03-01"
        assert _basic(structured, "expire_date")["value"] == "2027-02-28"

    def test_currency_explicit_then_inferred(self) -> None:
        explicit = extract_structured("币种：人民币；其他略")
        assert _basic(explicit, "currency")["status"] == "ok"
        inferred = extract_structured("服务费以人民币结算，共 12 万元。")
        cur = _basic(inferred, "currency")
        assert cur["status"] == "inferred" and cur["value"] == "人民币" and isinstance(cur["pos"], int)

    def test_title_fallback_from_md_heading(self) -> None:
        structured = extract_structured("# 客户数据处理服务协议\n\n甲方：某公司")
        title = _basic(structured, "contract_title")
        assert title["status"] == "inferred"
        assert title["value"].endswith("协议")

    def test_party_stops_before_credit_code_paren(self) -> None:
        text = "甲方（采购方）：XX科技有限公司（统一社会信用代码：91310000MA1FL8X20A）"
        party = _basic(extract_structured(text), "party_a")
        assert party["value"] == "XX科技有限公司"


class TestClauses:
    def test_presence_contains_snippet_and_pos(self) -> None:
        text = "第五条 违约责任：任何一方违约应支付违约金。"
        breach = extract_structured(text)["clauses"]["breach_clause"]
        assert breach["status"] == "present"
        assert breach["keywords_hit"] in ("违约", "责任")
        assert "违约金" in breach["snippet"]

    def test_absence_recorded_not_empty(self) -> None:
        """规范要求：缺失也必须入库，不允许只返回空结果。"""
        clauses = extract_structured("仅含付款条款：验收后支付。")["clauses"]
        ip = clauses["ip_clause"]
        assert ip["status"] == "absent" and ip["snippet"] is None and ip["pos"] is None
        assert set(clauses) == {
            "payment_clause", "delivery_clause", "acceptance_clause", "breach_clause",
            "confidential_clause", "data_clause", "ip_clause", "dispute_clause"}


class TestDocxRoundtrip:
    def test_docx_full_pipeline(self, tmp_path) -> None:
        """生成与 mock inst-001 同构的 docx → 全文提取 → 结构化断言。"""
        import docx

        path = tmp_path / "c.docx"
        document = docx.Document()
        document.add_heading("GPU 服务器集群采购合同", level=0)
        for para in (
            "合同编号：HT-2026-0301",
            "甲方（采购方）：XX科技有限公司（统一社会信用代码：91310000MA1FL8X20A）",
            "乙方（供应商）：华信计算设备有限公司",
            "合同总金额为人民币 1,860,000 元（大写：壹佰捌拾陆万元整），含税。",
            "预付款比例为合同总金额的 50%。",
            "剩余 50% 货款于到货后 90 个工作日内支付。",
            "因履行本合同发生争议的，任何一方可向甲方所在地人民法院提起诉讼。",
        ):
            document.add_paragraph(para)
        document.save(path)

        text, file_type = extract_raw_text(str(path))
        assert file_type == "docx"
        structured = extract_structured(text)
        assert _basic(structured, "contract_title")["status"] == "inferred"
        assert _basic(structured, "contract_no")["value"] == "HT-2026-0301"
        assert _basic(structured, "amount")["value"] == 1_860_000.0
        clauses = structured["clauses"]
        assert clauses["payment_clause"]["status"] == "present"
        assert clauses["acceptance_clause"]["status"] == "absent"
        assert clauses["breach_clause"]["status"] == "absent"


class TestErrorPaths:
    def test_unsupported_extension_raises(self, tmp_path) -> None:
        weird = tmp_path / "contract.exe"
        weird.write_bytes(b"MZ")
        with pytest.raises(ParseError):
            extract_raw_text(str(weird))

    def test_empty_docx_raises(self, tmp_path) -> None:
        import docx

        path = tmp_path / "empty.docx"
        docx.Document().save(path)
        with pytest.raises(ParseError):
            extract_raw_text(str(path))


@pytest.mark.skipif(shutil.which("tesseract") is None,
                    reason="宿主机未装 tesseract——OCR 真实路径由容器内探针覆盖")
class TestOcrGated:
    def test_png_ocr_roundtrip(self, tmp_path) -> None:
        from PIL import Image, ImageDraw

        path = tmp_path / "scan.png"
        img = Image.new("RGB", (600, 120), "white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 40), "HELLO CONTRACT 12345", fill="black")
        img.save(path)
        text, file_type = extract_raw_text(str(path))
        assert file_type == "png" and "CONTRACT" in text.upper()
