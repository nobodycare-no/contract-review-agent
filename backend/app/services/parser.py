"""合同文档解析器：docx/md/txt → 全文；pdf → pypdf；png/jpg → tesseract OCR。

结构化输出契约（SDD §8）：所有字段三元组 {value, pos, status}；
金额归一化为数值元(value)+raw_text；日期归一化 YYYY-MM-DD；缺失也入库(status=missing)。
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path

from app.core.config import get_settings


class ParseError(Exception):
    pass


# ---------- 全文提取 ----------

def extract_raw_text(file_path: str) -> tuple[str, str]:
    """按扩展名提取全文，返回 (text, file_type)。"""
    path = Path(file_path)
    ext = path.suffix.lower()
    data = path.read_bytes()

    if ext in (".md", ".txt"):
        for enc in ("utf-8", "gb18030"):
            try:
                return data.decode(enc), ext[1:]
            except UnicodeDecodeError:
                continue
        raise ParseError("文本编码无法识别")

    if ext == ".docx":
        return _extract_docx(data), "docx"

    if ext == ".pdf":
        return _extract_pdf(data), "pdf"

    if ext in (".png", ".jpg", ".jpeg"):
        return _extract_ocr(data), ext[1:]

    raise ParseError(f"不支持的附件类型: {ext}")


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ParseError(f"DOCX 解析失败: {e}") from e
    parts: list[str] = []
    for para in document.paragraphs:
        style = getattr(para.style, "name", "") or ""
        text = para.text.strip()
        if not text:
            continue
        parts.append(f"# {text}" if style.startswith("Heading") or style == "Title" else text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ParseError("DOCX 中未提取到任何文本")
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as e:
        raise ParseError(f"PDF 解析失败: {e}") from e
    text = "\n".join(p for p in pages if p.strip())
    if not text:
        raise ParseError("PDF 无文字层——疑似扫描件，OCR 流程请走图片路径")
    return text


def _extract_ocr(data: bytes) -> str:
    import pytesseract
    from PIL import Image

    settings = get_settings()
    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
    try:
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img, lang=settings.ocr_lang)
    except Exception as e:
        raise ParseError(f"OCR 识别失败: {e}") from e
    text = text.strip()
    if not text:
        raise ParseError("图片扫描件 OCR 后无有效文字（可能清晰度不足或非合同内容）")
    return text


# ---------- 归一化工具（SDD §8） ----------

_DATE_RE = re.compile(r"((?:19|20)\d{2})\s*[年.\-/]\s*(\d{1,2})\s*[月.\-/]\s*(\d{1,2})\s*日?")


def _norm_date(raw: str) -> str | None:
    m = _DATE_RE.search(raw)
    if not m:
        return None
    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    try:
        import datetime as _dt

        _dt.date(y, mo, d)
    except ValueError:
        return None
    return f"{y:04d}-{mo:02d}-{d:02d}"


def _norm_amount_yuan(digits: str, wan: bool) -> float:
    value = float(digits.replace(",", "").replace("，", ""))
    return value * 10_000 if wan else value


_CURRENCY_WORDS = ("人民币", "美元", "欧元", "日元", "港币")


# ---------- 结构化提取 ----------

BASIC_PATTERNS: dict[str, list[str]] = {
    "contract_title": [r"合同名称[：:]\s*(.+)"],
    "contract_no": [r"(?:合同编号|协议编号)[：:]\s*([A-Za-z0-9\-/]+)"],
    "party_a": [
        r"甲方[（(][^）)]*[）)]\s*[：:]\s*([^\n，,。;；（(]{2,50})",
        r"甲方\s*[：:]\s*([^\n，,。;；（(]{2,50})",
    ],
    "party_b": [
        r"乙方[（(][^）)]*[）)]\s*[：:]\s*([^\n，,。;；（(]{2,50})",
        r"乙方\s*[：:]\s*([^\n，,。;；（(]{2,50})",
    ],
    # 金额：捕获数字串与可选“万”，归一化为数值元
    "amount": [r"(?:合同总?金额|总价|总金额)[^0-9]{0,12}?([0-9][0-9,，]*(?:\.[0-9]+)?)\s*(万)?\s*元"],
    "currency": [r"(?:币种|货币)\s*[：:]\s*(\S{1,10})"],
    "effective_date": [r"(?:生效日期|生效时间)\s*[：:]?\s*((?:19|20)\d{2}\s*[年.\-/].{0,10})"],
    "expire_date": [r"(?:到期日期|到期时间|终止日期)\s*[：:]?\s*((?:19|20)\d{2}\s*[年.\-/].{0,10})"],
}

_TITLE_FALLBACK_RE = re.compile(r"^#{0,3}\s*\**(.{2,40}?(?:合同|协议))\**\s*$", re.M)

CLAUSE_KEYWORDS: dict[str, list[str]] = {
    "payment_clause": ["付款", "支付", "预付", "尾款", "结算"],
    "delivery_clause": ["交付", "交货", "供货", "工期"],
    "acceptance_clause": ["验收", "检验标准"],
    "breach_clause": ["违约", "赔偿", "责任"],
    "confidential_clause": ["保密", "机密"],
    "data_clause": ["个人信息", "数据安全", "数据处理"],
    "ip_clause": ["知识产权", "著作权", "成果归属"],
    "dispute_clause": ["争议", "纠纷", "管辖", "仲裁", "诉讼"],
}


def _field_entry(value, pos: int | None, status: str, **extra) -> dict:
    entry = {"value": value, "pos": pos, "status": status}
    entry.update(extra)
    return entry


def extract_structured(text: str) -> dict:
    """正则兜底提取基本信息与八类条款定位（LLM 增强在 services 层叠加）。"""
    basic: dict[str, dict] = {}
    for field, patterns in BASIC_PATTERNS.items():
        entry = _field_entry(None, None, "missing")
        for pat in patterns:
            m = re.search(pat, text)
            if not m:
                continue
            raw = m.group(1).strip()
            if field == "amount":
                raw_amount = re.sub(r"^[^0-9]+", "", m.group(0))  # 从原匹配截取，保留原始空格
                entry = _field_entry(
                    _norm_amount_yuan(m.group(1), bool(m.group(2))),
                    m.start(), "ok",
                    raw_text=raw_amount, unit="CNY")
            elif field in ("effective_date", "expire_date"):
                normalized = _norm_date(raw)
                if normalized:
                    entry = _field_entry(normalized, m.start(), "ok", raw_text=raw)
                continue
            else:
                entry = _field_entry(raw, m.start(), "ok")
            break
        basic[field] = entry

    # 标题回退：首个形如「……合同/协议」的行（含 md 标题），status=inferred
    if basic["contract_title"]["status"] == "missing":
        m = _TITLE_FALLBACK_RE.search(text)
        if m:
            basic["contract_title"] = _field_entry(m.group(1).strip(), m.start(1), "inferred")

    # 币种推断：全文出现货币词且未显式声明
    if basic["currency"]["status"] == "missing":
        idx = next((text.find(w) for w in _CURRENCY_WORDS if text.find(w) >= 0), -1)
        if idx >= 0:
            word = next(w for w in _CURRENCY_WORDS if text.find(w) == idx)
            basic["currency"] = _field_entry(word, idx, "inferred")

    clauses: dict[str, dict] = {}
    for name, keywords in CLAUSE_KEYWORDS.items():
        found = None
        for kw in keywords:
            idx = text.find(kw)
            if idx >= 0:
                snippet = text[idx:idx + 160].replace("\n", " ")
                found = {"keywords_hit": kw, "snippet": snippet,
                         "pos": idx, "status": "present"}
                break
        clauses[name] = found or {"keywords_hit": None, "snippet": None,
                                  "pos": None, "status": "absent"}

    return {"basic_info": basic, "clauses": clauses, "meta": {"char_len": len(text)}}


def safe_json_loads(raw: str | None, default):
    if not raw:
        return default
    try:
        return json.loads(raw)
    except Exception:
        return default
