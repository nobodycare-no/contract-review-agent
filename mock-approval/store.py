"""mock 审批系统内存注册表：审批单运行态 + 评论接收。reset 可复现演示。"""
from __future__ import annotations

import copy
import io
import threading

import docx
from PIL import Image, ImageDraw, ImageFont

from contracts_def import APPROVALS, DATA_PROCESSING_MD, SCAN_PAGE_LINES, TEMPLATES


class MockStore:
    """线程安全的仿真审批单仓库（内存态，进程重启即复位）。"""

    def __init__(self):
        self._lock = threading.Lock()
        self._comments: dict[str, list[dict]] = {}
        self._reset_locked()

    def _reset_locked(self):
        self._approvals = copy.deepcopy(APPROVALS)
        self._comments = {a["instance_id"]: [] for a in self._approvals}

    def reset(self):
        with self._lock:
            self._reset_locked()

    # ---------- 查询 ----------

    def list_pending(self, limit: int = 20) -> list[dict]:
        with self._lock:
            rows = [a for a in self._approvals if a["status"] == "pending"]
            return [{
                "instance_id": a["instance_id"],
                "approval_code": a["approval_code"],
                "approval_title": a["approval_title"],
                "applicant_name": a["applicant_name"],
                "apply_time": a["apply_time"],
                "attachment_count": len(a["attachments"]),
            } for a in rows[:limit]]

    def get_detail(self, instance_id: str) -> dict | None:
        with self._lock:
            a = next((x for x in self._approvals if x["instance_id"] == instance_id), None)
            if a is None:
                return None
            return {
                "instance_id": a["instance_id"],
                "approval_code": a["approval_code"],
                "approval_title": a["approval_title"],
                "applicant_name": a["applicant_name"],
                "apply_time": a["apply_time"],
                "form_data": a["form_data"],
                "attachments": [dict(att) for att in a["attachments"]],
                "status": a["status"],
            }

    # ---------- 附件渲染 ----------

    @staticmethod
    def render_attachment(template: str, file_name: str,
                          instance_id: str) -> tuple[bytes, str]:
        """按模板动态生成附件字节流，返回 (bytes, content_type)。"""
        lower = file_name.lower()
        if template == "data_processing_md" or lower.endswith((".md", ".txt")):
            text = DATA_PROCESSING_MD
            if template != "data_processing_md":
                text = f"# {file_name}\n（演示占位正文）\n"
            return text.encode("utf-8"), "text/markdown"

        if lower.endswith(".png"):
            img = _render_png(SCAN_PAGE_LINES)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue(), "image/png"

        if lower.endswith(".docx"):
            tpl = TEMPLATES.get(template)
            if tpl is None:
                raise KeyError(f"未知模板: {template}")
            document = docx.Document()
            document.add_heading(tpl["title"], level=0)
            for para in tpl["paragraphs"]:
                document.add_paragraph(para)
            buf = io.BytesIO()
            document.save(buf)
            return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        raise KeyError(f"无法渲染的附件类型: {file_name}")

    # ---------- 评论回写 ----------

    def add_comment(self, instance_id: str, comment_text: str) -> dict:
        with self._lock:
            a = next((x for x in self._approvals if x["instance_id"] == instance_id), None)
            if a is None:
                return {"write_status": "failed", "error": "实例不存在"}
            self._comments[instance_id].append({
                "comment_id": len(self._comments[instance_id]) + 1,
                "author": "AI-合同审查助手",
                "comment_text": comment_text,
            })
            return {"write_status": "success",
                    "comment_id": self._comments[instance_id][-1]["comment_id"]}

    def list_comments(self, instance_id: str) -> list[dict]:
        with self._lock:
            return list(self._comments.get(instance_id, []))


def _render_png(lines: list[str]) -> "Image.Image":
    from PIL import Image, ImageDraw, ImageFont

    width, line_h = 900, 56
    img = Image.new("RGB", (width, line_h * (len(lines) + 2)), "white")
    draw = ImageDraw.Draw(img)
    font = None
    for candidate in ("C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf",
                      "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"):
        try:
            font = ImageFont.truetype(candidate, 30)
            break
        except OSError:
            continue
    if font is None:
        font = ImageFont.load_default()

    y = line_h // 2
    for line in lines:
        draw.text((40, y), line, fill="black", font=font)
        y += line_h
    return img


# 模块级单例
store = MockStore()
